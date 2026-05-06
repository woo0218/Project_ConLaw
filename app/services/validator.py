import io
import re

from docx import Document
from fastapi import HTTPException

from app.config import REQUIRED_PLACEHOLDERS


FORM_PLACEHOLDERS: list[str] = [
    "{{법령근거}}",
    "{{조항번호}}",
    "{{현장상황}}",
    "{{감리의견}}",
    "{{시정요구사항}}",
    "{{종합의견}}",
    "{{기타사항}}",
]

OFFICIAL_FORM_MARKERS: tuple[str, ...] = (
    "감리보고서",
    "건축공사 감리세부기준",
    "공사감리자",
    "관계전문기술자",
    "확인 및 의견",
    "기타사항",
    "종합의견",
)


def collect_document_text(doc: Document) -> str:
    values: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            values.append(text)

    for table in doc.tables:
        for row in table.rows:
            seen_cells: set[int] = set()
            for cell in row.cells:
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                text = cell.text.strip()
                if text:
                    values.append(text)

    return "\n".join(values)


def looks_like_official_report_form(doc: Document) -> bool:
    text = collect_document_text(doc)
    if "감리보고" not in text:
        return False

    marker_count = 0
    for marker in OFFICIAL_FORM_MARKERS:
        if marker in text:
            marker_count += 1
    return marker_count >= 2


def validate_template(filename: str, file_bytes: bytes) -> list[str]:
    """
    Validate uploaded docx. Return ordered deduplicated placeholder list.
    Raises HTTPException on any failure.
    """
    if not filename.lower().endswith(".docx"):
        raise HTTPException(400, "docx 파일만 허용됩니다.")

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception:
        raise HTTPException(400, "파일이 손상되었거나 올바른 docx 형식이 아닙니다.")

    seen: list[str] = []
    seen_set: set[str] = set()

    def collect(text: str) -> None:
        for m in re.findall(r"\{\{[^}]+\}\}", text):
            if m not in seen_set:
                seen.append(m)
                seen_set.add(m)

    for para in doc.paragraphs:
        collect(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    collect(para.text)

    if not seen:
        if looks_like_official_report_form(doc):
            return FORM_PLACEHOLDERS
        raise HTTPException(400, "템플릿에서 입력 가능한 항목을 찾을 수 없습니다.")

    missing = [p for p in REQUIRED_PLACEHOLDERS if p not in seen_set]
    if missing:
        raise HTTPException(
            400,
            f"필수 항목 누락: {missing}. {{{{법령근거}}}}, {{{{조항번호}}}}는 반드시 포함되어야 합니다.",
        )

    return seen
