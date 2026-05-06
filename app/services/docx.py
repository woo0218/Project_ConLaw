import io
import logging
import re
from typing import Dict, List, Tuple

from docx import Document

from app.config import REQUIRED_PLACEHOLDERS
from app.services.validator import FORM_PLACEHOLDERS, looks_like_official_report_form


logger = logging.getLogger(__name__)


class DocxService:
    def _collect_placeholders_from_doc(self, doc: Document) -> list[str]:
        seen: list[str] = []
        seen_set: set[str] = set()

        def collect(text: str) -> None:
            for m in re.findall(r"\{\{[^}]+\}\}", text):
                if m not in seen_set:
                    seen.append(m)
                    seen_set.add(m)

        for para in doc.paragraphs:
            collect(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        collect(para.text)

        return seen

    def extract_placeholders(self, file_bytes: bytes) -> list[str]:
        doc = Document(io.BytesIO(file_bytes))
        return self._collect_placeholders_from_doc(doc)

    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> None:
        if not paragraph.runs:
            return

        text = "".join(run.text for run in paragraph.runs)
        if not text:
            return

        original_text = text
        for key, value in replacements.items():
            if key in text:
                text = text.replace(key, value)

        if text == original_text:
            return

        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)

    def _unique_row_cells(self, row) -> list:
        cells = []
        seen: set[int] = set()
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id in seen:
                continue
            seen.add(cell_id)
            cells.append(cell)
        return cells

    def _row_text(self, row) -> str:
        return " ".join(cell.text.strip() for cell in self._unique_row_cells(row) if cell.text.strip())

    def _find_rows_containing(self, doc: Document, labels: tuple[str, ...]) -> list:
        rows = []
        for table in doc.tables:
            for row in table.rows:
                row_text = self._row_text(row)
                if any(label in row_text for label in labels):
                    rows.append(row)
        return rows

    def _best_fill_cell(self, row, labels: tuple[str, ...]):
        cells = self._unique_row_cells(row)
        label_indexes = [
            idx
            for idx, cell in enumerate(cells)
            if any(label in cell.text for label in labels)
        ]
        start = max(label_indexes) + 1 if label_indexes else 0

        for cell in cells[start:]:
            if not cell.text.strip():
                return cell

        for cell in reversed(cells[start:] or cells):
            text = cell.text.strip()
            if not any(label in text for label in labels):
                return cell
        return None

    def _write_cell(self, cell, text: str, append: bool = True) -> None:
        clean_text = str(text).strip()
        if not cell or not clean_text:
            return

        current = cell.text.strip()
        if current and clean_text in current:
            return
        if current and append:
            cell.text = current + "\n" + clean_text
        elif not current:
            cell.text = clean_text

    def _write_near_label(self, doc: Document, labels: tuple[str, ...], text: str, append: bool = True) -> bool:
        for row in self._find_rows_containing(doc, labels):
            cell = self._best_fill_cell(row, labels)
            if cell is not None:
                self._write_cell(cell, text, append=append)
                return True
        return False

    def _build_form_sections(self, law: dict, llm_sections: dict) -> dict[str, str]:
        title = str(law.get("title") or "").strip()
        article_no = str(law.get("article_no") or "").strip()
        law_basis = " ".join(value for value in (title, article_no) if value)

        opinion = str(llm_sections.get("{{감리의견}}") or "").strip()
        correction = str(llm_sections.get("{{시정요구사항}}") or "").strip()
        summary = str(llm_sections.get("{{종합의견}}") or "").strip()
        etc = str(llm_sections.get("{{기타사항}}") or "").strip()

        if not summary:
            summary_parts = [part for part in (opinion, correction) if part]
            summary = "\n".join(summary_parts)
        if not summary:
            summary = (
                f"{law_basis}를 근거로 현장 시공 상태를 검토한 결과, "
                "구조 안전성 및 품질 확보를 위하여 지적 사항에 대한 시정 조치와 재확인이 필요합니다."
            )
        if not etc:
            etc = f"법령 근거: {law_basis}" if law_basis else ""

        sections: dict[str, str] = {
            "{{종합의견}}": summary,
            "{{기타사항}}": etc,
        }

        return sections

    def _fill_official_form(self, doc: Document, law: dict, llm_sections: dict) -> None:
        sections = self._build_form_sections(law, llm_sections)

        # Official report forms contain administrative/signature fields that must
        # remain untouched. Only fill the report narrative rows.
        if not self._write_near_label(doc, ("기타사항",), sections["{{기타사항}}"], append=False):
            logger.warning("기타사항 입력 대상 셀을 찾지 못했습니다.")

        if not self._write_near_label(doc, ("종합의견",), sections["{{종합의견}}"], append=False):
            logger.warning("종합의견 입력 대상 셀을 찾지 못했습니다.")

    def fill_template(self, file_bytes: bytes, law: dict, llm_sections: dict) -> Tuple[bytes, List[str]]:
        doc = Document(io.BytesIO(file_bytes))
        actual_placeholders = self._collect_placeholders_from_doc(doc)
        form_mode = not actual_placeholders and looks_like_official_report_form(doc)

        replacements: Dict[str, str] = {}

        title = law.get("title") or ""
        article_no = law.get("article_no") or ""

        if "{{법령근거}}" in REQUIRED_PLACEHOLDERS or True:
            replacements["{{법령근거}}"] = str(title)
        if "{{조항번호}}" in REQUIRED_PLACEHOLDERS or True:
            replacements["{{조항번호}}"] = str(article_no)

        for key, value in llm_sections.items():
            if not isinstance(value, (str, int, float)):
                continue
            replacements[str(key)] = str(value)

        if form_mode:
            self._fill_official_form(doc, law, llm_sections)
        else:
            for para in doc.paragraphs:
                self._replace_in_paragraph(para, replacements)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, replacements)

        remaining: list[str] = []
        remaining_set: set[str] = set()

        def collect_remaining(text: str) -> None:
            for m in re.findall(r"\{\{[^}]+\}\}", text):
                if m not in remaining_set:
                    remaining.append(m)
                    remaining_set.add(m)

        if not form_mode:
            for para in doc.paragraphs:
                collect_remaining(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            collect_remaining(para.text)

        unmatched: list[str] = []
        for placeholder in remaining:
            if placeholder not in replacements:
                unmatched.append(placeholder)

        if unmatched:
            logger.warning("템플릿 미매핑 placeholder: %s", unmatched)

        if unmatched:
            fill_map = {ph: "[미입력]" for ph in unmatched}

            for para in doc.paragraphs:
                self._replace_in_paragraph(para, fill_map)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_in_paragraph(para, fill_map)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue(), unmatched
